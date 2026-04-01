from __future__ import annotations

import sys
from pathlib import Path
from datetime import datetime

ROOT = Path(__file__).resolve().parent
COMMON = ROOT / "services" / "common"
if str(COMMON) not in sys.path:
    sys.path.insert(0, str(COMMON))

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=True)
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(md_path: Path, docx_path: Path):
    """Convert markdown comparison report to DOCX format"""
    
    # Read markdown content
    content = md_path.read_text(encoding='utf-8')
    lines = content.split('\n')
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Track current section
    in_table = False
    table_data = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines unless we're building something
        if not line:
            if not in_table:
                doc.add_paragraph()
            i += 1
            continue
        
        # Main Title (# )
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            title = doc.add_heading(text, level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.runs[0]
            run.font.color.rgb = RGBColor(214, 48, 49)  # Adobe red
            run.font.size = Pt(24)
            run.bold = True
        
        # Heading 2 (## )
        elif line.startswith('## '):
            text = line[3:].strip()
            h = doc.add_heading(text, level=1)
            run = h.runs[0]
            run.font.color.rgb = RGBColor(214, 48, 49)
            run.font.size = Pt(18)
        
        # Heading 3 (### )
        elif line.startswith('### '):
            text = line[4:].strip()
            h = doc.add_heading(text, level=2)
            run = h.runs[0]
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(50, 50, 50)
        
        # Heading 4 (#### )
        elif line.startswith('#### '):
            text = line[5:].strip()
            h = doc.add_heading(text, level=3)
            run = h.runs[0]
            run.font.size = Pt(12)
            run.bold = True
        
        # Bold text (**text**)
        elif line.startswith('**') and line.endswith('**'):
            text = line[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
        
        # Table detection
        elif '|' in line and line.startswith('|'):
            if not in_table:
                in_table = True
                table_data = []
            
            # Parse table row
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(cells)
            
            # Check if next line is separator or end of table
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if not next_line.startswith('|'):
                    # End of table, create it
                    create_table(doc, table_data)
                    in_table = False
                    table_data = []
                elif '---' in next_line:
                    # Skip separator line
                    i += 1
            else:
                # End of file, create table
                create_table(doc, table_data)
                in_table = False
        
        # Horizontal rule (---)
        elif line == '---':
            doc.add_paragraph('_' * 80)
        
        # Bullet points (- or *)
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            text = process_inline_formatting(text)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
        
        # Numbered list
        elif line[0].isdigit() and '. ' in line[:4]:
            text = line.split('. ', 1)[1].strip()
            text = process_inline_formatting(text)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
        
        # Regular paragraph
        else:
            text = process_inline_formatting(line)
            p = doc.add_paragraph()
            add_formatted_text(p, text)
        
        i += 1
    
    # Save document
    doc.save(str(docx_path))
    print(f"✅ DOCX created: {docx_path}")


def create_table(doc, table_data):
    """Create a formatted table in the document"""
    if not table_data:
        return
    
    num_rows = len(table_data)
    num_cols = len(table_data[0])
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    for row_idx, row_data in enumerate(table_data):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            
            # Clean up formatting markers
            text = cell_data.replace('**', '')
            cell.text = text
            
            # Format header row
            if row_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Set background color for header (shading)
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'D63031')  # Adobe red
                cell._element.get_or_add_tcPr().append(shading_elm)
            else:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)


def process_inline_formatting(text):
    """Process inline markdown formatting markers"""
    # We'll keep markers for now and process them in add_formatted_text
    return text


def add_formatted_text(paragraph, text):
    """Add text with inline formatting (bold, italic, links)"""
    import re
    
    # Find all formatting patterns
    # **bold**, _italic_, [text](url), `code`
    
    parts = []
    current_pos = 0
    
    # Pattern for bold, links, etc.
    patterns = [
        (r'\*\*(.+?)\*\*', 'bold'),
        (r'_(.+?)_', 'italic'),
        (r'\[(.+?)\]\((.+?)\)', 'link'),
        (r'`(.+?)`', 'code'),
    ]
    
    # Find all matches
    matches = []
    for pattern, format_type in patterns:
        for match in re.finditer(pattern, text):
            matches.append((match.start(), match.end(), match, format_type))
    
    # Sort by position
    matches.sort(key=lambda x: x[0])
    
    # Process text with formatting
    for start, end, match, format_type in matches:
        # Add text before match
        if current_pos < start:
            paragraph.add_run(text[current_pos:start])
        
        # Add formatted text
        if format_type == 'bold':
            run = paragraph.add_run(match.group(1))
            run.bold = True
        elif format_type == 'italic':
            run = paragraph.add_run(match.group(1))
            run.italic = True
        elif format_type == 'link':
            link_text = match.group(1)
            url = match.group(2)
            run = paragraph.add_run(f"{link_text}")
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.underline = True
        elif format_type == 'code':
            run = paragraph.add_run(match.group(1))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(199, 37, 78)
        
        current_pos = end
    
    # Add remaining text
    if current_pos < len(text):
        paragraph.add_run(text[current_pos:])


def main():
    md_path = Path("data/outputs/real_creators_comparison.md")
    docx_path = Path("data/outputs/real_creators_comparison.docx")
    
    if not md_path.exists():
        print(f"❌ Markdown file not found: {md_path}")
        return
    
    print(f"📄 Converting {md_path} to DOCX...")
    markdown_to_docx(md_path, docx_path)
    print(f"✅ Done! Document saved to: {docx_path}")


if __name__ == "__main__":
    main()
